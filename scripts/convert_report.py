import os
import re
import sys
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def parse_markdown(md_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    return lines

def clean_markdown_links(text):
    # Replaces [text](url) with just text for a clean docx representation
    return re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)

def clean_markdown_math(text):
    # Strip double dollar block math and simplify
    text = text.replace(r'$$\mathbf{u}_{\text{dynamic}} = \sigma(\alpha) \cdot \mathbf{v}_{\text{short}} + (1 - \sigma(\alpha)) \cdot \mathbf{v}_{\text{long}}$$', 
                        'u_dynamic = σ(α) · v_short + (1 - σ(α)) · v_long')
    text = text.replace(r'$$\hat{y}_{ui} = \mathbf{u}_{\text{dynamic}} \cdot \mathbf{q}_i + b_i$$', 
                        'y_ui = u_dynamic · q_i + b_i')
    
    # Strip inline math and replace with clean unicode
    replacements = {
        r'$N-2$': 'N-2',
        r'$N-1$': 'N-1',
        r'$N$': 'N',
        r'$\alpha$': 'α',
        r'$\sigma(\alpha)$': 'σ(α)',
        r'$b_i$': 'b_i',
        r'$\mathbf{v}_{\text{short}}$': 'v_short',
        r'$\mathbf{v}_{\text{long}}$': 'v_long',
        r'$\mathbf{u}_{\text{dynamic}}$': 'u_dynamic',
        r'$\mathbf{u}_{\text{dynamic}} \cdot \mathbf{q}_i + b_i$': 'u_dynamic · q_i + b_i',
        r'$\hat{y}_{ui} = \mathbf{u}_{\text{dynamic}} \cdot \mathbf{q}_i + b_i$': 'y_ui = u_dynamic · q_i + b_i',
        r'$\hat{y}_{ui} = \text{user\_repr} \cdot \mathbf{q}_i + b_i$': 'y_ui = user_repr · q_i + b_i',
        r'\sigma(\alpha)': 'σ(α)',
        r'$$\mathbf{u}_{\text{dynamic}} = \sigma(\alpha) \cdot \mathbf{v}_{\text{short}} + (1 - \sigma(\alpha)) \cdot \mathbf{v}_{\text{long}}$$': 'u_dynamic = σ(α) · v_short + (1 - σ(α)) · v_long',
        r'$$\hat{y}_{ui} = \mathbf{u}_{\text{dynamic}} \cdot \mathbf{q}_i + b_i$$': 'y_ui = u_dynamic · q_i + b_i',
        r'user\_repr': 'user_repr'
    }
    
    for k, v in replacements.items():
        text = text.replace(k, v)
        
    # General cleanup of remaining inline math patterns
    text = re.sub(r'\$(\\mathbf|\\math[a-zA-Z]+)?{?([^$]+)}?\$', r'\2', text)
    text = text.replace(r'\cdot', '·').replace(r'\sigma', 'σ').replace(r'\alpha', 'α').replace(r'\_', '_')
    text = text.replace('$', '') # remove any remaining dollar signs
    return text

def add_formatted_text(paragraph, text):
    # Handle GfM HTML/markdown links
    text = clean_markdown_links(text)
    # Handle math formulas
    text = clean_markdown_math(text)
    # Split by bold '**'
    parts = text.split('**')
    for idx, part in enumerate(parts):
        is_bold = (idx % 2 == 1)
        # Split by italic '*'
        subparts = part.split('*')
        for sidx, subpart in enumerate(subparts):
            is_italic = (sidx % 2 == 1)
            if not subpart:
                continue
            run = paragraph.add_run(subpart)
            if is_bold:
                run.bold = True
            if is_italic:
                run.italic = True

def add_custom_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Format header row
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = ""
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header.strip())
        run.bold = True
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(10)
        
    # Populate data rows
    for row in rows:
        row_cells = table.add_row().cells
        for idx, val in enumerate(row):
            if idx >= len(row_cells):
                break
            val_clean = val.strip()
            row_cells[idx].text = ""
            p = row_cells[idx].paragraphs[0]
            # Center alignment for metrics, left for models
            if idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Format text
            add_formatted_text(p, val_clean)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
                run.font.size = Pt(9.5)

def build_docx(md_path, docx_path):
    print(f"Reading from {md_path}")
    lines = parse_markdown(md_path)
    
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Set default style font
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(10.5)
    
    in_table = False
    table_headers = []
    table_rows = []
    
    in_code_block = False
    code_lines = []
    
    idx = 0
    while idx < len(lines):
        line = lines[idx]
        line_stripped = line.strip()
        
        # Code block handler
        if line_stripped.startswith("```"):
            if in_code_block:
                # End of code block
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                
                code_text = "".join(code_lines).rstrip()
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(80, 80, 80)
                
                in_code_block = False
                code_lines = []
            else:
                # Start of code block
                in_code_block = True
            idx += 1
            continue
            
        if in_code_block:
            code_lines.append(line)
            idx += 1
            continue
            
        # Table handler
        if line_stripped.startswith("|"):
            # It's a table row
            parts = [p.strip() for p in line_stripped.split("|")[1:-1]]
            
            # Check if it's separator row (e.g. |:---:|:---:|)
            is_sep = all(re.match(r'^:?-+:?$', p) for p in parts)
            
            if is_sep:
                # Skip separator
                idx += 1
                continue
                
            if not in_table:
                in_table = True
                table_headers = parts
                table_rows = []
            else:
                table_rows.append(parts)
            idx += 1
            continue
        else:
            if in_table:
                # Flush the table
                add_custom_table(doc, table_headers, table_rows)
                # Space after table
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(6)
                
                in_table = False
                table_headers = []
                table_rows = []
        
        if not line_stripped:
            idx += 1
            continue
            
        # Skip horizontal rules
        if line_stripped == "---":
            idx += 1
            continue
            
        # Headings
        if line_stripped.startswith("# "):
            heading_text = line_stripped[2:]
            h = doc.add_heading(level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = h.add_run(heading_text)
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(20)
            run.bold = True
            run.font.color.rgb = RGBColor(26, 82, 118) # Deep blue
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(12)
            
        elif line_stripped.startswith("## "):
            heading_text = line_stripped[3:]
            h = doc.add_heading(level=2)
            run = h.add_run(heading_text)
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(14)
            run.bold = True
            run.font.color.rgb = RGBColor(46, 134, 193) # Mid blue
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(6)
            
        elif line_stripped.startswith("### "):
            heading_text = line_stripped[4:]
            h = doc.add_heading(level=3)
            run = h.add_run(heading_text)
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor(52, 73, 94) # Dark gray-blue
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after = Pt(4)
            
        elif line_stripped.startswith("#### "):
            heading_text = line_stripped[5:]
            h = doc.add_heading(level=4)
            run = h.add_run(heading_text)
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(11)
            run.bold = True
            run.font.color.rgb = RGBColor(100, 110, 120) # Slate gray-blue
            h.paragraph_format.space_before = Pt(6)
            h.paragraph_format.space_after = Pt(3)
            
        # List items
        elif line_stripped.startswith("- ") or line_stripped.startswith("* "):
            item_text = line_stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_formatted_text(p, item_text)
            for r in p.runs:
                r.font.name = 'Microsoft YaHei'
                r.font.size = Pt(10.5)
                
        elif re.match(r'^\d+\.\s', line_stripped):
            match = re.match(r'^(\d+)\.\s(.*)', line_stripped)
            num = match.group(1)
            item_text = match.group(2)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            add_formatted_text(p, item_text)
            for r in p.runs:
                r.font.name = 'Microsoft YaHei'
                r.font.size = Pt(10.5)
                
        # Normal paragraphs
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            add_formatted_text(p, line_stripped)
            for r in p.runs:
                r.font.name = 'Microsoft YaHei'
                r.font.size = Pt(10.5)
                
        idx += 1
        
    # Final check for tables
    if in_table:
        add_custom_table(doc, table_headers, table_rows)
        
    print(f"Saving DOCX to {docx_path}")
    try:
        doc.save(docx_path)
        print("Success!")
    except PermissionError:
        fallback_path = docx_path.replace(".docx", "_已更新.docx")
        print(f"Warning: {docx_path} is currently open and locked by another program.")
        print(f"Saving to fallback path instead: {fallback_path}")
        doc.save(fallback_path)
        print("Success! (Saved to fallback file)")

if __name__ == "__main__":
    md = "实验报告.md"
    docx = "实验报告.docx"
    build_docx(md, docx)
